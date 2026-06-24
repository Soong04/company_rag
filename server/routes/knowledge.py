# ============================================================================
# 知识库 CRUD 路由
# ============================================================================
import os
import logging
import jieba
from concurrent.futures import ThreadPoolExecutor
from flask import Blueprint, request, jsonify, current_app, send_file
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from models import db
from models.knowledge import Category, Document

logger = logging.getLogger(__name__)

knowledge_bp = Blueprint('knowledge', __name__, url_prefix='/api/knowledge')

# ---------- 文档向量化后台处理 ----------
# 线程池：最多同时处理 2 个文档，避免 Ollama 过载
_doc_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix='doc_vect')


# ========== 分类管理 ==========

@knowledge_bp.route('/categories', methods=['GET'])
@login_required
def get_categories():
    """
    获取分类列表（树形结构）
    GET /api/knowledge/categories
    """
    # 获取所有分类
    all_categories = Category.query.order_by(Category.sort_order, Category.id).all()

    # 构建树形结构
    def build_tree(parent_id=0):
        tree = []
        for cat in all_categories:
            if cat.parent_id == parent_id:
                children = build_tree(cat.id)
                node = cat.to_dict()
                if children:
                    node['children'] = children
                tree.append(node)
        return tree

    tree = build_tree(0)

    return jsonify({
        'code': 200,
        'data': tree
    })


@knowledge_bp.route('/categories', methods=['POST'])
@login_required
def create_category():
    """
    创建分类
    POST /api/knowledge/categories
    """
    data = request.get_json()
    if not data or not data.get('name'):
        return jsonify({'code': 400, 'message': '分类名称不能为空'}), 400

    category = Category(
        name=data['name'],
        parent_id=data.get('parent_id', 0),
        sort_order=data.get('sort_order', 0),
        description=data.get('description', '')
    )
    db.session.add(category)
    db.session.commit()

    return jsonify({'code': 200, 'message': '创建成功', 'data': {'category': category.to_dict()}})


@knowledge_bp.route('/categories/<int:category_id>', methods=['PUT'])
@login_required
def update_category(category_id):
    """
    更新分类
    PUT /api/knowledge/categories/<id>
    """
    category = Category.query.get_or_404(category_id)
    data = request.get_json()
    if not data:
        return jsonify({'code': 400, 'message': '请求数据为空'}), 400

    if 'name' in data:
        category.name = data['name']
    if 'parent_id' in data:
        category.parent_id = data['parent_id']
    if 'sort_order' in data:
        category.sort_order = data['sort_order']
    if 'description' in data:
        category.description = data['description']

    db.session.commit()

    return jsonify({'code': 200, 'message': '更新成功', 'data': {'category': category.to_dict()}})


@knowledge_bp.route('/categories/<int:category_id>', methods=['DELETE'])
@login_required
def delete_category(category_id):
    """
    删除分类
    DELETE /api/knowledge/categories/<id>
    """
    category = Category.query.get_or_404(category_id)

    # 检查是否有子分类
    sub_count = Category.query.filter_by(parent_id=category_id).count()
    if sub_count > 0:
        return jsonify({'code': 400, 'message': '该分类下有子分类，请先删除子分类'}), 400

    # 检查是否有文档
    doc_count = Document.query.filter_by(category_id=category_id).count()
    if doc_count > 0:
        return jsonify({'code': 400, 'message': '该分类下有文档，请先删除文档'}), 400

    db.session.delete(category)
    db.session.commit()

    return jsonify({'code': 200, 'message': '删除成功'})


# ========== 文档管理 ==========

@knowledge_bp.route('/documents', methods=['GET'])
@login_required
def get_documents():
    """
    获取文档列表
    GET /api/knowledge/documents?page=1&size=20&category_id=&keyword=&status=
    """
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 20, type=int)
    category_id = request.args.get('category_id', type=int)
    keyword = request.args.get('keyword', '').strip()
    status = request.args.get('status', type=int)

    query = Document.query

    if category_id:
        query = query.filter_by(category_id=category_id)
    if keyword:
        query = query.filter(
            Document.title.like(f'%{keyword}%') |
            Document.content_summary.like(f'%{keyword}%')
        )
    if status is not None:
        query = query.filter_by(status=status)

    pagination = query.order_by(Document.create_time.desc()).paginate(
        page=page, per_page=size, error_out=False
    )

    documents = [doc.to_dict() for doc in pagination.items]

    return jsonify({
        'code': 200,
        'data': {
            'list': documents,
            'total': pagination.total,
            'page': page,
            'size': size
        }
    })


@knowledge_bp.route('/documents/<int:doc_id>', methods=['GET'])
@login_required
def get_document(doc_id):
    """
    获取文档详情
    GET /api/knowledge/documents/<id>
    """
    document = Document.query.get_or_404(doc_id)
    return jsonify({
        'code': 200,
        'data': {'document': document.to_dict()}
    })


@knowledge_bp.route('/documents/<int:doc_id>/preview', methods=['GET'])
@login_required
def preview_document(doc_id):
    """
    预览文档内容
    GET /api/knowledge/documents/<id>/preview
    返回文档的完整文本内容和元数据
    """
    document = Document.query.get_or_404(doc_id)
    return jsonify({
        'code': 200,
        'data': {
            'id': document.id,
            'title': document.title,
            'file_type': document.file_type,
            'file_size': document.file_size,
            'content': document.content or '',
            'category_name': document.category.name if document.category else None,
            'create_time': document.create_time.strftime('%Y-%m-%d %H:%M:%S'),
        }
    })


@knowledge_bp.route('/documents/<int:doc_id>/file', methods=['GET'])
@login_required
def download_document(doc_id):
    """
    获取文档原始文件（用于下载/浏览器预览）
    GET /api/knowledge/documents/<id>/file
    """
    document = Document.query.get_or_404(doc_id)
    if not document.file_path:
        return jsonify({'code': 404, 'message': '文件路径不存在'}), 404

    # 尝试多个可能的路径（兼容不同启动目录）
    stored_path = document.file_path.replace('/', os.sep).replace('\\', os.sep)
    possible_paths = [
        stored_path,
        os.path.abspath(stored_path),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', stored_path),
        os.path.join(os.getcwd(), stored_path),
    ]

    real_path = None
    for p in possible_paths:
        expanded = os.path.abspath(p)
        if os.path.exists(expanded):
            real_path = expanded
            break

    if not real_path:
        return jsonify({'code': 404, 'message': '文件不存在，请重新上传'}), 404

    # 根据文件类型设置 MIME
    mime_map = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'txt': 'text/plain; charset=utf-8',
        'md': 'text/markdown; charset=utf-8',
    }
    mime = mime_map.get(document.file_type, 'application/octet-stream')

    return send_file(
        real_path,
        mimetype=mime,
        as_attachment=False,
        download_name=document.title + '.' + document.file_type,
    )


@knowledge_bp.route('/documents', methods=['POST'])
@login_required
def upload_document():
    """
    上传文档
    POST /api/knowledge/documents
    FormData: file, category_id, title
    """
    if 'file' not in request.files:
        return jsonify({'code': 400, 'message': '请选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'code': 400, 'message': '文件名为空'}), 400

    category_id = request.form.get('category_id', type=int)
    if not category_id:
        return jsonify({'code': 400, 'message': '请选择分类'}), 400

    # 检查分类是否存在
    category = Category.query.get(category_id)
    if not category:
        return jsonify({'code': 400, 'message': '分类不存在'}), 400

    # 从原始文件名提取扩展名（secure_filename 会破坏中文文件名）
    original_name = file.filename or ''
    ext = original_name.rsplit('.', 1)[-1].lower() if '.' in original_name else ''
    if ext not in current_app.config.get('ALLOWED_EXTENSIONS', {'pdf', 'docx', 'txt', 'md'}):
        return jsonify({'code': 400, 'message': f'不支持的文件类型: .{ext}'}), 400

    # 保存文件（使用绝对路径）
    import time
    ts = int(time.time())
    upload_folder = os.path.abspath(current_app.config.get('UPLOAD_FOLDER', './uploads'))
    os.makedirs(upload_folder, exist_ok=True)

    filename = secure_filename(original_name) or f'document_{ts}'
    unique_name = f"{ts}_{filename}"
    file_path = os.path.join(upload_folder, unique_name)
    file.save(file_path)

    # 读取文件内容（含表格 → Markdown 转换）
    content = ''
    try:
        if ext == 'txt' or ext == 'md':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif ext == 'pdf':
            content = _extract_pdf_with_tables(file_path)
        elif ext == 'docx':
            content = _extract_docx_with_tables(file_path)
        elif ext == 'doc':
            content = _extract_doc(file_path)
    except Exception as e:
        return jsonify({'code': 500, 'message': f'文件解析失败: {str(e)}'}), 500

    # 限制存储内容大小（避免 MySQL TEXT 字段溢出，完整内容可从原始文件读取）
    MAX_DB_CONTENT = 100000
    if len(content) > MAX_DB_CONTENT:
        content = content[:MAX_DB_CONTENT] + '\n\n[内容已截断，完整内容请查看原始文件]'

    title = request.form.get('title') or original_name or filename
    file_size = os.path.getsize(file_path)

    # 创建文档记录（状态：处理中）
    doc_id_for_thread = None
    document = Document(
        category_id=category_id,
        title=title,
        file_type=ext,
        file_path=file_path,
        file_size=file_size,
        content_summary=content[:200] + '...' if len(content) > 200 else content,
        content=content,  # 存储完整内容用于全文搜索
        status=1,  # 处理中
        upload_user_id=current_user.id
    )
    db.session.add(document)
    db.session.commit()
    doc_id_for_thread = document.id

    # ---------- 后台异步进行文本分块和向量化 ----------
    def _process_document_async(doc_id, doc_content, doc_title, flask_app):
        """在线程池中处理文档向量化（重试 1 次）"""
        with flask_app.app_context():
            from services.vector_service import VectorService
            from services.deepseek_service import DeepSeekService
            from services.hybrid_retriever import HybridRetriever
            from models import db, Document

            doc = db.session.get(Document, doc_id)
            if not doc:
                logger.warning(f'[向量化] 文档 #{doc_id} 已被删除，跳过')
                return

            cfg = flask_app.config
            persist_dir = cfg.get('CHROMA_PERSIST_DIR', './chroma_data')
            collection_name = cfg.get('CHROMA_COLLECTION_NAME', 'enterprise_knowledge')
            deepseek_api_key = cfg.get('DEEPSEEK_API_KEY', '')
            deepseek_base_url = cfg.get('DEEPSEEK_BASE_URL', 'https://api.deepseek.com')
            embedding_model = cfg.get('DEEPSEEK_EMBEDDING_MODEL', 'deepseek-embedding')

            def _do_vectorize():
                """实际的向量化逻辑（可供重试）"""
                chunks = chunk_text(doc_content)
                doc.chunk_count = len(chunks)

                vector_service = VectorService(
                    persist_dir=persist_dir,
                    collection_name=collection_name
                )
                llm_service = DeepSeekService(
                    api_key=deepseek_api_key,
                    base_url=deepseek_base_url,
                    embedding_model=embedding_model
                )

                # 生成向量并存入 Chroma
                embeddings = llm_service.embed_documents(chunks)
                metadatas = [
                    {'doc_id': str(doc.id), 'title': doc.title, 'chunk_index': i}
                    for i in range(len(chunks))
                ]
                vector_service.add_documents(chunks, metadatas, embeddings)

                # 更新 BM25 索引
                try:
                    hybrid = HybridRetriever(persist_dir=persist_dir)
                    hybrid.add_documents(chunks, metadatas)
                except Exception as bm25_e:
                    logger.warning(f'[向量化] BM25 索引更新失败（不影响主流程）: {bm25_e}')

            try:
                _do_vectorize()
                doc.status = 2  # 已完成
                db.session.commit()
                logger.info(f'[向量化] 文档 #{doc_id}「{doc_title}」向量化完成（{doc.chunk_count} 个分块）')

            except Exception as e:
                db.session.rollback()
                logger.error(f'[向量化] 文档 #{doc_id}「{doc_title}」首次失败: {e}')

                # 自动重试一次（应对 Ollama 临时故障）
                try:
                    logger.info(f'[向量化] 文档 #{doc_id} 开始重试...')
                    _do_vectorize()
                    doc.status = 2
                    db.session.commit()
                    logger.info(f'[向量化] 文档 #{doc_id} 重试成功')
                except Exception as e2:
                    db.session.rollback()
                    doc.status = -1  # 失败
                    db.session.commit()
                    logger.error(f'[向量化] 文档 #{doc_id} 重试仍失败: {e2}')

    # 提交到线程池（非阻塞）
    _doc_executor.submit(
        _process_document_async,
        doc_id_for_thread, content, title,
        current_app._get_current_object()
    )

    return jsonify({
        'code': 200,
        'message': '上传成功',
        'data': {'document': document.to_dict()}
    })


@knowledge_bp.route('/search', methods=['GET'])
@login_required
def search_documents():
    """
    全文搜索文档内容
    GET /api/knowledge/search?q=关键词&page=1&size=20
    """
    keyword = request.args.get('q', '').strip()
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 20, type=int)

    if not keyword:
        return jsonify({'code': 400, 'message': '请输入搜索关键词'}), 400

    # 使用 FULLTEXT 索引搜索（MATCH AGAINST）+ LIKE 降级兜底
    # MySQL 需要 ngram parser 支持中文，默认 ngram_token_size=2
    from sqlalchemy import or_, text, func

    try:
        # 尝试 FULLTEXT 搜索（BOOLEAN MODE 支持 +* 前缀匹配）
        # 对每个关键词追加 * 以实现前缀匹配（如 attend* -> attendance）
        ft_keyword = ' '.join([f'+{w}*' for w in keyword.split() if w])
        if not ft_keyword:
            ft_keyword = f'+{keyword}*'

        match_condition = or_(
            func.match(Document.title, ft_keyword, text("IN BOOLEAN MODE")),
            func.match(Document.content, ft_keyword, text("IN BOOLEAN MODE")),
            Document.content_summary.like(f'%{keyword}%'),
        )
        query = Document.query.filter(match_condition)
        # 按全文相关度排序
        relevance = func.match(Document.title, ft_keyword, text("IN BOOLEAN MODE")) + \
                    func.match(Document.content, ft_keyword, text("IN BOOLEAN MODE"))
        pagination = query.order_by(relevance.desc(), Document.create_time.desc()).paginate(
            page=page, per_page=size, error_out=False
        )
    except Exception:
        # FULLTEXT 不可用时降级为 LIKE 搜索
        query = Document.query.filter(
            Document.title.like(f'%{keyword}%') |
            Document.content_summary.like(f'%{keyword}%') |
            Document.content.like(f'%{keyword}%')
        )
        pagination = query.order_by(Document.create_time.desc()).paginate(
            page=page, per_page=size, error_out=False
        )

    documents = [doc.to_search_dict(keyword) for doc in pagination.items]

    return jsonify({
        'code': 200,
        'data': {
            'list': documents,
            'total': pagination.total,
            'page': page,
            'size': size,
            'keyword': keyword,
        }
    })


@knowledge_bp.route('/documents/<int:doc_id>', methods=['PUT'])
@login_required
def update_document(doc_id):
    """
    更新文档信息
    PUT /api/knowledge/documents/<id>
    """
    document = Document.query.get_or_404(doc_id)
    data = request.get_json()
    if not data:
        return jsonify({'code': 400, 'message': '请求数据为空'}), 400

    if 'title' in data:
        document.title = data['title']
    if 'category_id' in data:
        document.category_id = data['category_id']
    if 'content_summary' in data:
        document.content_summary = data['content_summary']

    db.session.commit()

    return jsonify({'code': 200, 'message': '更新成功', 'data': {'document': document.to_dict()}})


@knowledge_bp.route('/documents/<int:doc_id>', methods=['DELETE'])
@login_required
def delete_document(doc_id):
    """
    删除文档
    DELETE /api/knowledge/documents/<id>
    """
    document = Document.query.get_or_404(doc_id)
    doc_id_str = str(document.id)
    logger.info(f'[删除] 开始删除文档 #{doc_id}「{document.title}」')

    # 从 Chroma 中删除向量
    try:
        from app import get_vector_service
        vector_service = get_vector_service()
        if vector_service:
            before = vector_service.get_collection_stats().get('total_documents', 0)
            vector_service.delete_by_metadata('doc_id', doc_id_str)
            after = vector_service.get_collection_stats().get('total_documents', 0)
            removed = before - after
            logger.info(f'[删除] Chroma 清理了 {removed} 个向量块（{before} → {after}）')
        else:
            logger.warning('[删除] Chroma 服务不可用，跳过向量清理')
    except Exception as e:
        logger.error(f'[删除] Chroma 向量删除失败: {e}')

    # 从 BM25 索引中删除
    try:
        from services.hybrid_retriever import HybridRetriever
        hybrid = HybridRetriever(
            persist_dir=current_app.config.get('CHROMA_PERSIST_DIR', './chroma_data')
        )
        hybrid.delete_by_metadata('doc_id', doc_id_str)
        logger.info(f'[删除] BM25 索引已清理')
    except Exception as e:
        logger.error(f'[删除] BM25 索引删除失败: {e}')

    # 删除文件
    if document.file_path and os.path.exists(document.file_path):
        try:
            os.remove(document.file_path)
            logger.info(f'[删除] 文件已删除: {document.file_path}')
        except Exception as e:
            logger.warning(f'[删除] 文件删除失败: {e}')

    db.session.delete(document)
    db.session.commit()
    logger.info(f'[删除] 文档 #{doc_id} 删除完成')

    return jsonify({'code': 200, 'message': '删除成功'})


@knowledge_bp.route('/documents/<int:doc_id>/revectorize', methods=['POST'])
@login_required
def revectorize_document(doc_id):
    """
    重新向量化文档（重试失败的文档）
    POST /api/knowledge/documents/<id>/revectorize
    """
    document = Document.query.get_or_404(doc_id)

    if document.status != -1:
        return jsonify({'code': 400, 'message': '只有处理失败的文档可以重新向量化'}), 400

    if not document.content:
        return jsonify({'code': 400, 'message': '文档内容为空，请重新上传'}), 400

    # 重置状态为处理中
    document.status = 1
    db.session.commit()

    # 提交到线程池重新处理
    _doc_executor.submit(
        _process_document_async,
        doc_id, document.content, document.title,
        current_app._get_current_object()
    )

    return jsonify({
        'code': 200,
        'message': '已重新提交向量化任务',
        'data': {'document': document.to_dict()}
    })


def _table_to_text(table, table_index: int) -> str:
    """将 DOCX 表格对象转换为自然语言描述"""
    rows_data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_data.append(cells)
    if not rows_data:
        return ''

    lines = [f'[表格 {table_index}]']
    header = rows_data[0]

    if len(rows_data) > 1:
        for row in rows_data[1:]:
            parts = []
            for i, cell in enumerate(row):
                col_name = header[i] if i < len(header) and header[i] else f'列{i+1}'
                parts.append(f'{col_name}：{cell}')
            if parts:
                lines.append('- ' + '；'.join(parts))
    else:
        lines.append('- ' + '；'.join(header))

    return '\n'.join(lines)


def _extract_docx_with_tables(filepath: str) -> str:
    """提取 DOCX 内容（段落 + 表格→自然语言，保持顺序）"""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(filepath)
    body = doc.element.body
    content = ''
    table_index = 0

    # 按文档顺序遍历：段落和表格交错出现
    para_idx = 0
    tbl_idx = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p' and para_idx < len(doc.paragraphs):
            text = doc.paragraphs[para_idx].text.strip()
            if text:
                content += text + '\n'
            para_idx += 1
        elif tag == 'tbl' and tbl_idx < len(doc.tables):
            table_index += 1
            content += _table_to_text(doc.tables[tbl_idx], table_index) + '\n'
            tbl_idx += 1

    return content


def _extract_doc(filepath: str) -> str:
    """Extract text from legacy .doc files (olefile fallback to python-docx)"""
    try:
        import olefile
        ole = olefile.OleFileIO(filepath)
        if ole.exists('WordDocument'):
            data = ole.openstream('WordDocument').read()
            text = ''.join(chr(b) for i, b in enumerate(data) if i % 2 == 0 and (32 <= b < 127 or b in (10, 13)))
            ole.close()
            if len(text.strip()) > 50:
                return text.strip()
        ole.close()
    except ImportError:
        pass
    except Exception:
        pass

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            return '\n'.join(paragraphs)
    except Exception:
        pass

    return '[File format is .doc, please convert to .docx and retry]'


def _extract_pdf_with_tables(filepath: str) -> str:
    """提取 PDF 内容（优先 pdfplumber 获取表格，降级 pypdf）"""
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            content = ''
            table_index = 0
            for page_num, page in enumerate(pdf.pages, 1):
                # 提取表格
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        if not table or not any(any(cell.strip() for cell in row if cell) for row in table):
                            continue
                        table_index += 1
                        lines = [f'[表格 {table_index}]']
                        header = [c.strip() if c else '' for c in (table[0] or [])]
                        for row in table[1:]:
                            if not row or all(not (c or '').strip() for c in row):
                                continue
                            parts = []
                            for i, cell in enumerate(row):
                                col_name = header[i] if i < len(header) and header[i] else f'列{i+1}'
                                parts.append(f'{col_name}：{(cell or "").strip()}')
                            if parts:
                                lines.append('- ' + '；'.join(parts))
                        content += '\n'.join(lines) + '\n'

                # 提取普通文本
                text = page.extract_text()
                if text:
                    content += text + '\n'
            return content
    except ImportError:
        pass
    except Exception:
        pass

    # 降级：pypdf 纯文本提取
    from pypdf import PdfReader
    reader = PdfReader(filepath)
    content = ''
    for page in reader.pages:
        content += page.extract_text() + '\n'
    return content


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list:
    """
    对文本进行分块处理
    使用 jieba 分词辅助，按句子切分再合并到指定大小
    :param text: 原始文本
    :param chunk_size: 每块最大字符数
    :param overlap: 块间重叠字符数
    :return: 文本块列表
    """
    import re
    # 按句号、问号、感叹号、换行分割句子
    sentences = re.split(r'([。！？\n])', text)
    # 重组句子（保留分隔符）
    sentence_list = []
    for i in range(0, len(sentences) - 1, 2):
        combined = sentences[i] + (sentences[i + 1] if i + 1 < len(sentences) else '')
        if combined.strip():
            sentence_list.append(combined.strip())
    if len(sentences) % 2 == 1 and sentences[-1].strip():
        sentence_list.append(sentences[-1].strip())

    if not sentence_list:
        # 如果没有分割出句子，则按字符直接切分
        chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            chunks.append(text[i:i + chunk_size])
        return chunks if chunks else [text]

    chunks = []
    current_chunk = ''
    for sentence in sentence_list:
        if len(current_chunk) + len(sentence) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # 保留 overlap 部分
            overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else ''
            current_chunk = overlap_text + sentence
        else:
            current_chunk += sentence

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks if chunks else [text]
